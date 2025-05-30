from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging
from .models import DocumentGenerationRequest, ChapterData, ReferenceData

logger = logging.getLogger(__name__)

def create_research_document(data: DocumentGenerationRequest, output_path: str) -> str:
    """
    Generates a Word document from the provided research data and saves it.
    Returns the name of the generated file.
    """
    try:
        doc = Document()

        # --- Basic Setup (can be expanded with formatting_options) ---
        # Set default font for the document (more robustly done by modifying styles)
        style = doc.styles['Normal']
        font = style.font
        font.name = data.formatting_options.get("font_family", 'Times New Roman')
        font.size = Pt(data.formatting_options.get("font_size_main", 12))

        # Line spacing (example for normal style)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = data.formatting_options.get("line_spacing", 1.5) # 1.5 lines

        # --- Title Page (Very Basic) ---
        doc.add_heading(data.research_title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph() # Spacer
        doc.add_paragraph(f"By: {data.student_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Specialization: {data.specialization}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Institution: {data.university_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # --- Table of Contents (Placeholder - python-docx doesn't auto-generate fully dynamic ToC easily) ---
        # You might need to instruct users to "Update Field" in Word.
        # Or use more advanced techniques or libraries if a fully automated ToC is critical for MVP.
        # For a very simple placeholder:
        # toc_heading = doc.add_heading('Table of Contents', level=1)
        # toc_paragraph = doc.add_paragraph()
        # run = toc_paragraph.add_run()
        # fldChar = OxmlElement('w:fldChar')
        # fldChar.set(qn('w:fldCharType'), 'begin')
        # run._r.append(fldChar)
        # run = toc_paragraph.add_run()
        # instrText = OxmlElement('w:instrText')
        # instrText.set(qn('xml:space'), 'preserve')
        # instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # Word field code for ToC
        # run._r.append(instrText)
        # run = toc_paragraph.add_run()
        # fldChar = OxmlElement('w:fldChar')
        # fldChar.set(qn('w:fldCharType'), 'end')
        # run._r.append(fldChar)
        # doc.add_page_break()
        logger.info("Skipping Table of Contents generation for MVP simplicity.")


        # --- Chapters ---
        # Define heading styles (can be done once)
        # styles = doc.styles
        # if not any(s.name == 'CustomHeading1' for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH):
        #     heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        #     heading1_style.base_style = styles['Heading 1']
        #     heading1_style.font.name = 'Arial' # Example
        #     heading1_style.font.size = Pt(16)


        for chapter in data.chapters:
            logger.info(f"Adding chapter: {chapter.title}")
            doc.add_heading(chapter.title, level=1) # Use built-in Heading 1
            # Split content into paragraphs. Assume content might have newlines.
            paragraphs = chapter.content.split('\n')
            for para_text in paragraphs:
                if para_text.strip(): # Add paragraph if not empty
                    doc.add_paragraph(para_text.strip())
            doc.add_paragraph() # Spacer after chapter content

        # --- References Section (Basic APA style example) ---
        if data.references:
            logger.info("Adding References section")
            doc.add_heading('References', level=1)
            # Sort references alphabetically if needed (complex for full APA)
            for ref in data.references:
                if ref.citation_apa:
                    # Add hanging indent for references (common in APA)
                    p = doc.add_paragraph(style='List Paragraph') # Or a custom reference style
                    p.paragraph_format.left_indent = Inches(0.0)
                    p.paragraph_format.first_line_indent = Inches(-0.5) # Negative for hanging
                    p.add_run(ref.citation_apa)
                else:
                    # Fallback if only partial data
                    doc.add_paragraph(f"Reference data missing for a source.", style='List Paragraph')


        file_name = f"project_{data.project_id}_{data.research_title.replace(' ', '_')[:30]}.docx"
        full_output_path = f"{output_path}/{file_name}"
        doc.save(full_output_path)
        logger.info(f"Document saved to {full_output_path}")
        return file_name, full_output_path

    except Exception as e:
        logger.error(f"Error generating document for project {data.project_id}: {e}", exc_info=True)
        raise # Re-raise the exception to be caught by the FastAPI handler